import logging

import requests
import base64
import json
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
import os
from logging_config import setup_logging

organization = "STMN-Group"
project = "Data capturing Solutions ART"
team = "SG1"
#Enter your personal ADO token in user env variable, \> setx AZURE_DEVOPS_PAT "your_token_here"
pat = os.getenv("AZURE_DEVOPS_PAT")
planid = 2935   # Use Sirios Master library ID
token = base64.b64encode(f":{pat}".encode()).decode()
headers = {
    "Authorization": f"Basic {token}",
    "Content-Type": "application/json"
}

class Rccia:

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.logger.debug("Rccia initialized")
        self.test_case_xref = None

        # self.logger = logging.getLogger(self.__class__.__name__)
        # self.logger.setLevel(logging.DEBUG)
        self.wit_wi_rel_tc = []
        self.wit_wi_tested_by = []

        # Prevent duplicate handlers if multiple objects are created
        # if not self.logger.handlers:
        #     # Create file handler
        #     file_handler = logging.FileHandler("robot.log", mode="w")
        #
        #     # Define log format
        #     formatter = logging.Formatter(
        #         "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        #     )
        #
        #     file_handler.setFormatter(formatter)
        #
        #     # Add handler to logger
        #     self.logger.addHandler(file_handler)



    def check_azure_devops_auth(self):
        pat = os.getenv("AZURE_DEVOPS_PAT")
        print(pat)
        if not pat:
            raise Exception("PAT not found in environment variables")

        auth = base64.b64encode(f":{pat}".encode()).decode()
        url = f"https://dev.azure.com/{organization}/_apis/projects?api-version=7.1"
        headers = {
            "Authorization": f"Basic {auth}"
        }

        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            print("Authentication successful")
            return True
        elif response.status_code == 401:
            raise Exception("Authentication failed: Invalid PAT")
        else:
            raise Exception(f"Unexpected error: {response.status_code} - {response.text}")

    def get_test_suite_doc_number(self, test_suite_id):
        """
        Get info about ADO test suite and get the text within the field Description, it also clean the text
        gather fron ADO from all HTML tag and formating.
        :param test_suite_id: test case number
        :return: clean text of desciption field
        """
        url= (f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems/{test_suite_id}?api-version=7.0")
        resp = requests.get(url, headers=headers)
        json_response = resp.json()
        try:
            test_spec_eqt_number = (BeautifulSoup(json_response["fields"]["System.Description"], "html.parser").
                                    get_text(separator=" ", strip=True))
            return test_spec_eqt_number
        except Exception as e:
            print(f"Test suite {test_suite_id} doesn't have proper description")
            print(f"The description is: {json_response["fields"]["System.Description"]}")
            print(e)
            return "incorrect or empty description"

    def get_test_suite_id_listing(self, ver_or_val):
        """
        This is the first step to create the xref list. This function get from API get all test suite under the
        test plan 'planid'. For each test suite it check if the test suite is under the test suite name passed
        in parameter (usually verification or validation
        :param ver_or_val: string verification or validation
        :return: Only the last test suite level
        """
        # Get the list of all test suite and expand all sub level test suite under test plan
        url = (f"https://dev.azure.com/{organization}/{project}/_apis/testplan/Plans/{planid}/suites?expand={{expand}}&api-version=7.0")
        resp = requests.get(url, headers=headers)
        json_response = resp.json()
        test_suite_list = []
        parent_test_suite_list = []
        leaf_test_suite = []
        for ts in json_response["value"]:
            self.logger.info(f"Create xref for {ts["id"]} - {ts["name"]}")
            try:
                if ts["parentSuite"]["name"] == ver_or_val:  # Check level 1 test suite if parent is under the parameter
                    test_suite_list.append(ts["id"])         # add the test suite id to a list
            except:
                continue
            # Create a list of test suite that have child test suite
            try:
                if ts["parentSuite"]["id"]:
                    parent_test_suite_list.append(ts["parentSuite"]["id"])
            except:
                continue

        for ts in json_response["value"]:
            try:
                if ts["parentSuite"]["id"] in test_suite_list:
                    leaf_test_suite.append(ts["id"])
            except:
                continue
        for ts in test_suite_list:
            if ts not in parent_test_suite_list:
                leaf_test_suite.append(ts)
        return leaf_test_suite

    def create_xref_test_case_id_corresponding_etq_doc_number(self, list_of_test_suite):
        """
        Create a class dictionary with indice: test case id and value : table [test suite document number and title
                                                        Word Document test case ID TCxxxxx]
        :param list_of_test_suite: Table, list of test suite id
        """
        xref = {}
        print("Processing: ", end="")
        for test_suite in list_of_test_suite:
            print(".", end="")
            self.logger.info(f"Processing test case ID corresponding ETQ of test suite {test_suite}")
            test_case_list = self.get_test_case_id_list2(test_suite)
            for test_case in test_case_list:
                xref[test_case] = [self.get_test_suite_doc_number(test_suite),
                                                  self.get_test_case_TCid(test_case)]
        print("")
        return xref

    def get_test_case_id_list2(self, test_suite):
        """
        Return WI id for the test case
        """
        url = (f"https://dev.azure.com/{organization}/{project}/_apis/testplan/Plans/{planid}/Suites/{test_suite}/TestCase?api-version=7.0")
        resp = requests.get(url, headers=headers)
        json_response = resp.json()
        test_case_list_jsonobj = json_response["value"]
        test_case_list = []
        pattern = r'\D\D\d{5,7}'
        prog = re.compile(pattern)
        for tc in test_case_list_jsonobj:
            test_case_list.append(tc["workItem"]["id"])
        return test_case_list

    def get_test_case_TCid(self, test_case):
        """
        Get form the test case title field the part that contain TCxxxx id
        :param test_case: int, test case id
        :return: String, only the TC id in format TCxxxxx
        """
        pattern = r'\D\D\d{5,7}'
        prog = re.compile(pattern)
        json_data = self.get_workitem_details(test_case)
        return prog.search(json_data["fields"]["System.Title"]).group(0)

    def get_wiql_query_result(self):
        url = (f"https://dev.azure.com/{organization}/{project}/{team}/_apis/wit/wiql?api-version=7.0")
        body = {
            "query": """
            SELECT      [System.Id], [System.Title],  [System.WorkItemType],  [System.State]      
            FROM WorkItems    
            WHERE  (  [System.WorkItemType] IN ('Story', 'Bug')            
            AND [Custom.ProductVersion] = '1.0.8'            
            AND [System.State] in ('Done','Close'        ))"""
                }

        resp = requests.post(url, headers=headers, json=body)
        resp.raise_for_status()
        data = resp.json()
        self.logger.info(f"Get the query result")
        workitem_id = []
        for item in data["workItems"]:
            # return only bug or story
            if item["id"] != None:
                workitem_id.append(item["id"])
        self.logger.info(f"The WIQL query found following WI list {workitem_id}")
        return workitem_id

    def get_workitem_details(self, workitem):
        self.logger.info(f"Processing WIT {workitem}")
        url = (
            f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems/{workitem}?$expand=Relations&api-version=7.0")
        resp = requests.get(url, headers=headers)
        return resp.json()
    def get_tested_by(self, jsondata):
        """
        From the json data it extract the list of tested by test case, so it ignore all other WI relation (child, test, ...)
        :param jsondata:
        :return: list
        """
        self.logger.info(f"Initiate tested by for {jsondata["id"]}")
        test_by_TC_id_list = []
        try:
            for relation in jsondata["relations"]:
                if relation["rel"] == "Microsoft.VSTS.Common.TestedBy-Forward":
                    workitem_id = relation["url"].split("/")[-1]  # Get the ADO workitem ID from json object
                    workitem_detail = self.get_workitem_details(workitem_id)
                    # use previoulsy found WI ID to get detailed info if the WI have a tag verification
                    # it will be added to the tested by list
                    if workitem_detail["fields"]["System.Tags"].lower() == "verification":
                        test_by_TC_id_list.append(int(workitem_id))
                        self.logger.info(f"The WIT {jsondata["id"]} have tested by")
                        if jsondata["id"] not in self.wit_wi_tested_by:
                            self.wit_wi_tested_by.append(jsondata["id"])
                    else:
                        # print(f"test case {workitem_id} is not verification")
                        pass
        except Exception as e:
            print(e)
        return test_by_TC_id_list

    def get_impact_analysis(self, jsondata):

        try:
            text = BeautifulSoup(jsondata["fields"]["Custom.ImpactAnalysis"], "html.parser").get_text(
            separator=" ", strip=True)

            self.log_wit_with_rel_tc(text, int(jsondata["id"]))

            return text
        except:
            return "This Field is empty in ADO"

    def log_wit_with_rel_tc(self, text, wit_id):
        pattern = r"RelatedTC_Start(.*?)RelatedTC_Stop"
        c_pattern = [r"TC\d{5}", r"UC\d{5}", r"DEV-\d{6-7}"]


        match = re.search(pattern, text, re.DOTALL)

        if match:
            extracted_text = match.group(1).strip()
            for patt in c_pattern:
                tc_match = re.search(patt, extracted_text)
                if tc_match:
                    if wit_id not in self.wit_wi_rel_tc:
                        self.wit_wi_rel_tc.append(wit_id)



    def get_change_log(self, jsondata):
        try:
            return BeautifulSoup(jsondata["fields"]["Custom.ChangelogEntry_fullText"], "html.parser").get_text(
            separator=" ", strip=True)

        except:
            return "This Field is empty in ADO"


    def get_title(self, jsondata):

        try:
            return BeautifulSoup(jsondata["fields"]["System.Title"], "html.parser").get_text(
            separator=" ", strip=True)

        except:
            return "This Field is empty in ADO"

    def get_tc_corresponding_spec_etq_number(self, test_case_id):
        try:
            doc_number = self.test_case_xref[str(test_case_id)][0]
        except KeyError:
            doc_number = f"Fail to retrieved {test_case_id}"
            self.logger.error(f"Fail to retrieved {test_case_id}")
        return doc_number

    def formating_affecting_doc(self, tested_by_list):
        """
        From ADO WI id extracted by tested_by this function get the related Test spec (DEV-xxxxx - title)
        And also add the related TCxxxx test case numbering instead of ADO id.
        :param tested_by_list:
        :return:
        """
        doc_list = {}
        for tc in tested_by_list:
            try:
                tc_list = doc_list[self.get_tc_corresponding_spec_etq_number(tc)]
                tc_list.append(self.test_case_xref[str(tc)][1])
                doc_list[self.get_tc_corresponding_spec_etq_number(tc)] = tc_list

            except:
                doc_list[self.get_tc_corresponding_spec_etq_number(tc)] = [self.test_case_xref[str(tc)][1]]
        text = []
        for doc, tcs in doc_list.items():
            tcs.sort()
            text.append(f"{doc} - {", ".join(tcs)}")
        return "\n".join(text)



if __name__ == "__main__":
    setup_logging()
    logger = logging.getLogger(__name__)
    x = Rccia()
    logger.info("Start MAIN")
    x.check_azure_devops_auth()
    # Step 1 - Create cross-reference between TC ADO work item ID to a corresponding Test spec ETQ number
    test_suite_list = x.get_test_suite_id_listing("Verification")
    create_xref = False   # True or false

    if create_xref:
        with open("xref.json", "w") as file:
            json.dump(x.create_xref_test_case_id_corresponding_etq_doc_number(test_suite_list),
                      file, indent=4)

    with open("xref.json", mode='r') as file:
        x.test_case_xref = json.load(file)

    # Step 2 - Get the list of all story and bug in the release
    iteration_workitem_list = x.get_wiql_query_result()
    # step 3 - From the list of workitems get the
    # get json data for the WI
    # WI title and write it in the file
    # WI change description and write it in the file
    # WI Impact analysis and write it in the file
    # WI Testby test case list ADO id
    # Associated test spec ETQ doc number of each WI related TC

    with open("export_rccia.csv", mode='w', encoding='UTF-8') as f:
        f.write(f"Issue tracker,Title,Change description,Impact analysis,Affected documents\n")
        rccia_table = [["Issue tracker", "Title", "Change Description", "Impact analysis", "Affected documents"]]
        for wi in iteration_workitem_list:
            print(f"Prcessing WI: {wi}")
            rccia_line = []
            wi_json_data = x.get_workitem_details(wi)
            f.write(f"\"{wi}\",")
            rccia_line.append(str(wi))
            f.write(f"\"{x.get_title(wi_json_data)}\",")
            rccia_line.append(x.get_title(wi_json_data))
            f.write(f"\"{x.get_change_log(wi_json_data)}\",")
            rccia_line.append(x.get_change_log(wi_json_data))
            f.write(f"\"{x.get_impact_analysis(wi_json_data)}\",")
            rccia_line.append(x.get_impact_analysis(wi_json_data))
            testedby_list = x.get_tested_by(wi_json_data)
            f.write(f"\"{x.formating_affecting_doc(testedby_list)}\"\n")
            rccia_line.append(x.formating_affecting_doc(testedby_list))
            rccia_table.append(rccia_line)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    rows = len(rccia_table)
    cols = len(rccia_table[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rccia_table):
        for j, value in enumerate(row):
            table.cell(i,j).text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    print("This is the list of WIT that have TC in related TC but not link tested By")
    for wit in x.wit_wi_rel_tc:
        if wit not in x.wit_wi_tested_by:
            print(f"{wit}, ", end="")
    doc.save("export_rccia.docx")







